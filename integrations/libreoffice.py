import asyncio
import logging
import os
import subprocess
import sys
import tempfile
from typing import Any, Dict, List, Optional

import openpyxl
import pandas as pd
from docx import Document

logger = logging.getLogger("LibreOffice")


class LibreOffice:
    def __init__(self, libreoffice_path: Optional[str] = None):
        if libreoffice_path:
            self.libreoffice_path = libreoffice_path
        else:
            if sys.platform == "win32":
                self.libreoffice_path = "soffice.exe"
            else:
                self.libreoffice_path = "libreoffice"
        self.temp_dir = tempfile.gettempdir()

    async def convert(self, input_file: str, output_format: str = 'pdf') -> Dict:
        try:
            cmd = [
                self.libreoffice_path, '--headless', '--convert-to', output_format,
                '--outdir', os.path.dirname(input_file) or self.temp_dir,
                input_file
            ]
            loop = asyncio.get_running_loop()
            result = await loop.run_in_executor(
                None,
                lambda: subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            )
            if result.returncode == 0:
                base = os.path.splitext(os.path.basename(input_file))[0]
                output_file = os.path.join(
                    os.path.dirname(input_file) or self.temp_dir,
                    f"{base}.{output_format}"
                )
                return {"success": True, "output_file": output_file}
            return {"success": False, "error": result.stderr}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def create_spreadsheet(self, data: List[List[Any]], filename: str) -> Dict:
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            for row in data:
                ws.append(row)
            wb.save(filename)
            return {"success": True, "file": filename}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def create_document(self, content: List[Dict], filename: str) -> Dict:
        try:
            doc = Document()
            for item in content:
                if item['type'] == 'heading':
                    doc.add_heading(item['text'], level=item.get('level', 1))
                elif item['type'] == 'paragraph':
                    doc.add_paragraph(item['text'])
            doc.save(filename)
            return {"success": True, "file": filename}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_spreadsheet(self, filename: str) -> List[List]:
        if filename.endswith('.csv'):
            df = pd.read_csv(filename)
        else:
            df = pd.read_excel(filename)
        return df.values.tolist()